import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Read the markdown draft
with open('submission/HACKATHON_REPORT_DRAFT.md', 'r', encoding='utf-8') as f:
    md = f.read()

# Helper: Add a heading

def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level).add_run(text)
    p.font.name = 'Calibri'
    p.font.size = Pt(16 if level==1 else 14 if level==2 else 12)
    return p

# Helper: Add a table from markdown

def add_table_from_md(doc, md_table):
    lines = [l.strip() for l in md_table.strip().split('\n') if l.strip()]
    if len(lines) < 2:
        return
    headers = [h.strip() for h in lines[0].split('|') if h.strip()]
    rows = [
        [c.strip() for c in row.split('|') if c.strip()]
        for row in lines[2:]
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for row in rows:
        row_cells = table.add_row().cells
        for i, c in enumerate(row):
            row_cells[i].text = c
    return table

# Parse the markdown draft into sections
sections = {}
current = None
for line in md.splitlines():
    if line.startswith('# '):
        current = line[2:].strip()
        sections[current] = []
    elif line.startswith('## '):
        current = line[3:].strip()
        sections[current] = []
    elif current:
        sections[current].append(line)

# Create the DOCX
report = Document()
report.add_heading('Hackathon Segmentation Report', 0)

for sec, content in sections.items():
    if sec.lower().startswith('appendix'):
        report.add_page_break()
    if sec.lower().startswith('title'):
        add_heading(report, '1. Title & Summary', 1)
        for l in content:
            if l.strip():
                report.add_paragraph(l.strip())
        continue
    if sec.lower().startswith('methodology'):
        add_heading(report, '2. Methodology', 1)
        for l in content:
            report.add_paragraph(l.strip())
        continue
    if sec.lower().startswith('results'):
        add_heading(report, '3. Results & Performance', 1)
        in_table = False
        table_lines = []
        for l in content:
            if l.strip().startswith('|'):
                in_table = True
                table_lines.append(l)
            elif in_table and not l.strip():
                add_table_from_md(report, '\n'.join(table_lines))
                in_table = False
                table_lines = []
            elif in_table:
                table_lines.append(l)
            else:
                report.add_paragraph(l.strip())
        if table_lines:
            add_table_from_md(report, '\n'.join(table_lines))
        continue
    if sec.lower().startswith('challenges'):
        add_heading(report, '4. Challenges & Solutions', 1)
        for l in content:
            report.add_paragraph(l.strip())
        continue
    if sec.lower().startswith('optimizations'):
        add_heading(report, '5. Optimizations', 1)
        for l in content:
            report.add_paragraph(l.strip())
        continue
    if sec.lower().startswith('failure'):
        add_heading(report, '6. Failure Case Analysis', 1)
        in_table = False
        table_lines = []
        for l in content:
            if l.strip().startswith('|'):
                in_table = True
                table_lines.append(l)
            elif in_table and not l.strip():
                add_table_from_md(report, '\n'.join(table_lines))
                in_table = False
                table_lines = []
            elif in_table:
                table_lines.append(l)
            else:
                report.add_paragraph(l.strip())
        if table_lines:
            add_table_from_md(report, '\n'.join(table_lines))
        continue
    if sec.lower().startswith('conclusion'):
        add_heading(report, '7. Conclusion & Future Work', 1)
        for l in content:
            report.add_paragraph(l.strip())
        continue
    if sec.lower().startswith('appendix'):
        add_heading(report, '8. Appendix', 1)
        for l in content:
            report.add_paragraph(l.strip())
        continue
    # fallback
    add_heading(report, sec, 2)
    for l in content:
        report.add_paragraph(l.strip())

# Save
report.save('submission/HACKATHON_REPORT.docx')
print('Report generated: submission/HACKATHON_REPORT.docx')



